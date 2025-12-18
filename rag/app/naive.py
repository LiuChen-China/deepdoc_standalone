#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import sys, os; __name__ == "__main__" and sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
import logging
import re
import os
from functools import reduce
from io import BytesIO
from timeit import default_timer as timer
from docx import Document
from docx.image.exceptions import InvalidImageStreamError, UnexpectedEndOfFileError, UnrecognizedImageError
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.opc.oxml import parse_xml
from markdown import markdown
from PIL import Image
try:
    from common.token_utils import num_tokens_from_string
    from rag.utils.file_utils import extract_embed_file, extract_links_from_pdf, extract_links_from_docx, extract_html
    from deepdoc.parser import DocxParser, ExcelParser, HtmlParser, JsonParser, MarkdownElementExtractor, MarkdownParser, PdfParser, TxtParser
    from deepdoc.parser.pdf_parser import PlainParser
    from rag.nlp import concat_img, find_codec, naive_merge, naive_merge_with_images, naive_merge_docx, rag_tokenizer, tokenize_chunks, tokenize_chunks_with_images, tokenize_table
except:
    from deepdoc_standalone.common.token_utils import num_tokens_from_string
    from deepdoc_standalone.rag.utils.file_utils import extract_embed_file, extract_links_from_pdf, extract_links_from_docx, extract_html
    from deepdoc_standalone.deepdoc.parser import DocxParser, ExcelParser, HtmlParser, JsonParser, MarkdownElementExtractor, MarkdownParser, PdfParser, TxtParser
    from deepdoc_standalone.deepdoc.parser.pdf_parser import PlainParser
    from deepdoc_standalone.rag.nlp import concat_img, find_codec, naive_merge, naive_merge_with_images, naive_merge_docx, rag_tokenizer, tokenize_chunks, tokenize_chunks_with_images, tokenize_table

def by_deepdoc(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, pdf_cls = None ,**kwargs):
    callback = callback
    binary = binary
    pdf_parser = pdf_cls() if pdf_cls else Pdf()
    sections, tables = pdf_parser(
        filename if not binary else binary,
        from_page=from_page,
        to_page=to_page,
        callback=callback
    )

    return sections, tables, pdf_parser

def by_plaintext(filename, binary=None, from_page=0, to_page=100000, callback=None, **kwargs):
    pdf_parser = PlainParser()
    sections, tables = pdf_parser(
        filename if not binary else binary,
        from_page=from_page,
        to_page=to_page,
        callback=callback
    )
    return sections, tables, pdf_parser



PARSERS = {
    "deepdoc":  by_deepdoc,
    "plaintext": by_plaintext,
}


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        imgs = paragraph._element.xpath('.//pic:pic')
        if not imgs:
            return None
        res_img = None
        for img in imgs:
            embed = img.xpath('.//a:blip/@r:embed')
            if not embed:
                continue
            embed = embed[0]
            try:
                related_part = document.part.related_parts[embed]
                image_blob = related_part.image.blob
            except UnrecognizedImageError:
                logging.info("Unrecognized image format. Skipping image.")
                continue
            except UnexpectedEndOfFileError:
                logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
                continue
            except InvalidImageStreamError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except UnicodeDecodeError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except Exception:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            try:
                image = Image.open(BytesIO(image_blob)).convert('RGB')
                if res_img is None:
                    res_img = image
                else:
                    res_img = concat_img(res_img, image)
            except Exception:
                continue

        return res_img

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __get_nearest_title(self, table_index, filename):
        """Get the hierarchical title structure before the table"""
        import re
        from docx.text.paragraph import Paragraph

        titles = []
        blocks = []

        # Get document name from filename parameter
        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"

        # Collect all document blocks while maintaining document order
        try:
            # Iterate through all paragraphs and tables in document order
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith('p'):  # Paragraph
                    p = Paragraph(block, self.doc)
                    blocks.append(('p', i, p))
                elif block.tag.endswith('tbl'):  # Table
                    blocks.append(('t', i, None))  # Table object will be retrieved later
        except Exception as e:
            logging.error(f"Error collecting blocks: {e}")
            return ""

        # Find the target table position
        target_table_pos = -1
        table_count = 0
        for i, (block_type, pos, _) in enumerate(blocks):
            if block_type == 't':
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1

        if target_table_pos == -1:
            return ""  # Target table not found

        # Find the nearest heading paragraph in reverse order
        nearest_title = None
        for i in range(len(blocks)-1, -1, -1):
            block_type, pos, block = blocks[i]
            if pos >= target_table_pos:  # Skip blocks after the table
                continue

            if block_type != 'p':
                continue

            if block.style and block.style.name and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:  # Support up to 7 heading levels
                            title_text = block.text.strip()
                            if title_text:  # Avoid empty titles
                                nearest_title = (level, title_text)
                                break
                except Exception as e:
                    logging.error(f"Error parsing heading level: {e}")

        if nearest_title:
            # Add current title
            titles.append(nearest_title)
            current_level = nearest_title[0]

            # Find all parent headings, allowing cross-level search
            while current_level > 1:
                found = False
                for i in range(len(blocks)-1, -1, -1):
                    block_type, pos, block = blocks[i]
                    if pos >= target_table_pos:  # Skip blocks after the table
                        continue

                    if block_type != 'p':
                        continue

                    if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                # Find any heading with a higher level
                                if level < current_level:
                                    title_text = block.text.strip()
                                    if title_text:  # Avoid empty titles
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as e:
                            logging.error(f"Error parsing parent heading: {e}")

                if not found:  # Break if no parent heading is found
                    break

            # Sort by level (ascending, from highest to lowest)
            titles.sort(key=lambda x: x[0])
            # Organize titles (from highest to lowest)
            hierarchy = [doc_name] + [t[1] for t in titles]
            return " > ".join(hierarchy)

        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        """
        解析DOCX文档，提取文本内容、图片和表格信息
        
        参数:
            filename: 文档文件路径
            binary: 文档二进制数据，若提供则优先使用
            from_page: 开始解析的页码（包含）
            to_page: 结束解析的页码（不包含）
        
        返回:
            tuple: (文本行列表, 表格列表)
                - 文本行列表: [(文本内容, 关联图片), ...]
                - 表格列表: [((None, 表格HTML), ""), ...]
        """
        # 初始化文档对象，根据是否提供二进制数据选择不同的初始化方式
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        
        # 页码计数器
        pn = 0
        # 存储解析的文本行和图片信息
        lines = []
        # 暂存上一个段落的图片
        last_image = None
        
        # 遍历文档中的所有段落
        for p in self.doc.paragraphs:
            # 如果超过指定的结束页码，退出循环
            if pn > to_page:
                break
            
            # 仅处理在指定页码范围内的内容
            if from_page <= pn < to_page:
                # 处理非空段落
                if p.text.strip():
                    # 特殊处理图片说明段落
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        # 尝试从最近的非说明段落获取图片
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        # 添加说明文本和关联的图片
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        # 处理普通段落，提取段落中的图片
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        # 如果有上一个未处理的图片，添加到当前图片列表前面
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        # 添加段落文本和图片列表
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    # 处理空段落，检查是否包含图片
                    if current_image := self.get_picture(self.doc, p):
                        # 如果已有处理的行，将图片添加到最近一行
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            # 否则暂存图片
                            last_image = current_image
            
            # 检查段落中的页码标记，更新页码计数器
            for run in p.runs:
                # 检查是否包含页面结束标记
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                # 检查是否包含分页符
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        
        # 处理每一行的图片列表，将多张图片合并为一张
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        # 解析文档中的表格
        tbls = []
        for i, tb in enumerate(self.doc.tables):
            # 获取表格附近的标题信息
            title = self.__get_nearest_title(i, filename)
            # 构建表格HTML
            html = "<table>"
            if title:
                html += f"<caption>Table Location: {title}</caption>"
            
            # 遍历表格的每一行
            for r in tb.rows:
                html += "<tr>"
                i = 0
                try:
                    # 处理单元格，支持合并相同内容的单元格
                    while i < len(r.cells):
                        span = 1
                        c = r.cells[i]
                        # 查找连续的相同内容单元格
                        for j in range(i + 1, len(r.cells)):
                            if c.text == r.cells[j].text:
                                span += 1
                                i = j
                            else:
                                break
                        i += 1
                        # 根据合并情况生成对应的HTML单元格
                        html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                except Exception as e:
                    logging.warning(f"Error parsing table, ignore: {e}")
                html += "</tr>"
            html += "</table>"
            # 添加表格到结果列表
            tbls.append(((None, html), ""))
        
        # 返回解析结果
        return new_line, tbls


    def to_markdown(self, filename=None, binary=None, inline_images: bool = True):
        """
        This function uses mammoth, licensed under the BSD 2-Clause License.
        """

        import base64
        import uuid

        import mammoth
        from markdownify import markdownify

        docx_file = BytesIO(binary) if binary else open(filename, "rb")

        def _convert_image_to_base64(image):
            try:
                with image.open() as image_file:
                    image_bytes = image_file.read()
                encoded = base64.b64encode(image_bytes).decode("utf-8")
                base64_url = f"data:{image.content_type};base64,{encoded}"

                alt_name = "image"
                alt_name = f"img_{uuid.uuid4().hex[:8]}"

                return {"src": base64_url, "alt": alt_name}
            except Exception as e:
                logging.warning(f"Failed to convert image to base64: {e}")
                return {"src": "", "alt": "image"}

        try:
            if inline_images:
                result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.img_element(_convert_image_to_base64))
            else:
                result = mammoth.convert_to_html(docx_file)

            html = result.value

            markdown_text = markdownify(html)
            return markdown_text

        finally:
            if not binary:
                docx_file.close()


class Pdf(PdfParser):
    def __init__(self):
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None, separate_tables_figures=False):
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        #logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge(zoomin=zoomin)
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))

        if separate_tables_figures:
            tbls, figures = self._extract_table_figure(True, zoomin, True, True, True)
            self._concat_downward()
            #logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls, figures
        else:
            tbls = self._extract_table_figure(True, zoomin, True, True)
            self._naive_vertical_merge()
            self._concat_downward()
            self._final_reading_order_merge()
            # self._filter_forpages()
            #logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def md_to_html(self, sections):
        if not sections:
            return []
        if isinstance(sections, type("")):
            text = sections
        elif isinstance(sections[0], type("")):
            text = sections[0]
        else:
            return []

        from bs4 import BeautifulSoup
        html_content = markdown(text)
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup

    def get_hyperlink_urls(self, soup):
        if soup:
            return set([a.get('href') for a in soup.find_all('a') if a.get('href')])
        return []

    def extract_image_urls_with_lines(self, text):
        md_img_re = re.compile(r"!\[[^\]]*\]\(([^)\s]+)")
        html_img_re = re.compile(r'src=["\\\']([^"\\\'>\\s]+)', re.IGNORECASE)
        urls = []
        seen = set()
        lines = text.splitlines()
        for idx, line in enumerate(lines):
            for url in md_img_re.findall(line):
                if (url, idx) not in seen:
                    urls.append({"url": url, "line": idx})
                    seen.add((url, idx))
            for url in html_img_re.findall(line):
                if (url, idx) not in seen:
                    urls.append({"url": url, "line": idx})
                    seen.add((url, idx))

        # cross-line
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(text, 'html.parser')
            newline_offsets = [m.start() for m in re.finditer(r"\n", text)] + [len(text)]
            for img_tag in soup.find_all('img'):
                src = img_tag.get('src')
                if not src:
                    continue

                tag_str = str(img_tag)
                pos = text.find(tag_str)
                if pos == -1:
                    # fallback
                    pos = max(text.find(src), 0)
                line_no = 0
                for i, off in enumerate(newline_offsets):
                    if pos <= off:
                        line_no = i
                        break
                if (src, line_no) not in seen:
                    urls.append({"url": src, "line": line_no})
                    seen.add((src, line_no))
        except Exception:
            pass

        return urls

    def load_images_from_urls(self, urls, cache=None):
        import requests
        from pathlib import Path

        cache = cache or {}
        images = []
        for url in urls:
            if url in cache:
                if cache[url]:
                    images.append(cache[url])
                continue
            img_obj = None
            try:
                #关闭下载
                if 0:
                    if url.startswith(('http://', 'https://')):
                        response = requests.get(url, stream=True, timeout=30)
                        if response.status_code == 200 and response.headers.get('Content-Type', '').startswith('image/'):
                            img_obj = Image.open(BytesIO(response.content)).convert('RGB')
                    else:
                        local_path = Path(url)
                        if local_path.exists():
                            img_obj = Image.open(url).convert('RGB')
                        else:
                            logging.warning(f"Local image file not found: {url}")
            except Exception as e:
                logging.error(f"Failed to download/open image from {url}: {e}")
            cache[url] = img_obj
            if img_obj:
                images.append(img_obj)
        return images, cache

    def __call__(self, filename, binary=None, separate_tables=True, delimiter=None, return_section_images=False):
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()

        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n', separate_tables=separate_tables)
        # To eliminate duplicate tables in chunking result, uncomment code below and set separate_tables to True in line 410.
        # extractor = MarkdownElementExtractor(remainder)
        extractor = MarkdownElementExtractor(txt)
        image_refs = self.extract_image_urls_with_lines(txt)
        element_sections = extractor.extract_elements(delimiter, include_meta=True)

        sections = []
        section_images = []
        image_cache = {}
        for element in element_sections:
            content = element["content"]
            start_line = element["start_line"]
            end_line = element["end_line"]
            urls_in_section = [ref["url"] for ref in image_refs if start_line <= ref["line"] <= end_line]
            imgs = []
            if urls_in_section:
                imgs, image_cache = self.load_images_from_urls(urls_in_section, image_cache)
            combined_image = None
            if imgs:
                combined_image = reduce(concat_img, imgs) if len(imgs) > 1 else imgs[0]
            sections.append((content, ""))
            section_images.append(combined_image)

        tbls = []
        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        if return_section_images:
            return sections, tbls, section_images
        return sections, tbls

def load_from_xml_v2(baseURI, rels_item_xml):
    """
    Return |_SerializedRelationships| instance loaded with the
    relationships contained in *rels_item_xml*. Returns an empty
    collection if *rels_item_xml* is |None|.
    """
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ('../NULL', 'NULL'):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels



def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
    对指定文件进行分块处理，支持 docx, pdf, excel, txt 等格式。
    功能：
        1. 根据语言判断是否为英文，决定后续分词策略。
        2. 解析文件内容，提取文本、表格、图片信息。
        3. 将文本切分为小块（chunks），每块的 token 数量不超过 parser_config 中的限制。
        4. 支持可视化增强（使用 LLM 的 IMAGE2TEXT 模型增强表格/图像识别）。
        5. 可选只返回章节级别分块。
    
    参数：
        filename: 文件路径或名称。
        binary: 文件二进制内容（可选）。
        from_page/to_page: 针对 pdf 文件的页码范围。
        lang: 语言标记，主要用于判断分词方式（中文/英文）。
        callback: 进度回调，用来报告进度，比如 `"Start to parse."`。
        kwargs: 额外配置，比如 `parser_config`、`tenant_id`、`section_only`。   
    """
    # 初始化变量：用于存储超链接和处理结果
    urls = set()  # 存储从文档中提取的超链接
    url_res = []  # 存储从超链接处理得到的分块结果
    
    # 如果没有提供二进制内容，则从文件路径读取
    if binary is None:
        with open(filename, "rb") as f:
            binary = f.read()

    # 判断文档语言是否为英文，用于后续分词策略选择
    is_english = lang.lower() == "english"  # is_english(cks)
    
    # 获取解析配置，设置默认值
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 512,  # 每块最大token数量
            "delimiter": "\n!?。；！？",  # 文本分割符
            "layout_recognize": kwargs.get("layout_recognize", "DeepDOC"),  # 布局识别方式
            "analyze_hyperlink": False  # 是否分析超链接
        }
    )
    # 初始化文档信息和结果存储
    final_sections = False  # 标记是否使用最终格式化的段落
    doc = {  # 文档元信息
        "file_name": filename,  # 文档名称
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))  # 标题分词
    }
    # 对标题进行细粒度分词
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    
    # 初始化结果列表和解析器引用
    res = []  # 存储最终分块结果
    pdf_parser = None  # PDF解析器引用
    section_images = None  # 存储段落关联的图片

    # 检查是否为根调用（非递归调用）
    is_root = kwargs.get("is_root", True)
    embed_res = []  # 存储嵌入文件的处理结果
    
    # 仅在根调用时处理嵌入文件
    if is_root:
        # 提取嵌入内容
        embeds = []
        if binary is not None:
            embeds = extract_embed_file(binary)  # 提取嵌入内容 就是类似docx里插入的附件
        else:
            raise Exception("Embedding extraction from file path is not supported.")

        # 递归处理每个嵌入文件并收集结果
        for embed_filename, embed_bytes in embeds:
            try:
                sub_res = chunk(embed_filename, binary=embed_bytes, lang=lang, callback=callback, is_root=False, **kwargs) or []
                embed_res.extend(sub_res)
            except Exception as e:
                # if callback:
                #     callback(0.05, f"Failed to chunk embed {embed_filename}: {e}")
                continue

    # 根据文件类型选择不同的处理逻辑
    # 处理DOCX文件
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")  # 报告解析开始
        
        # 如果需要分析超链接且是根调用
        if parser_config.get("analyze_hyperlink", False) and is_root:
            urls = extract_links_from_docx(binary)
            for index, url in enumerate(urls):
                html_bytes, metadata = extract_html(url)
                if not html_bytes:
                    continue
                try:
                    sub_url_res = chunk(url, html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
                except Exception as e:
                    logging.info(f"Failed to chunk url in registered file type {url}: {e}")
                    sub_url_res = chunk(f"{index}.html", html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
                url_res.extend(sub_url_res)

        # 修复docx解析中的bug："There is no item named 'word/NULL' in the archive"
        _SerializedRelationships.load_from_xml = load_from_xml_v2
        
        # 解析docx文件内容
        sections, tables = Docx()(filename, binary)

        # 处理表格内容
        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")  # 报告解析完成

        # 记录合并开始时间
        st = timer()

        # 合并docx内容为分块
        chunks, images = naive_merge_docx(
            sections, int(parser_config.get("chunk_token_num", 128)), 
            parser_config.get("delimiter", "\n!?。；！？")
        )

        # 如果只需返回段落级别的分块
        if kwargs.get("section_only", False):
            chunks.extend(embed_res)
            chunks.extend(url_res)
            return chunks

        # 处理带图像的分块
        res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images))
        #logging.info("naive_merge({}): {}".format(filename, timer() - st))
        
        # 添加嵌入文件和超链接的处理结果
        res.extend(embed_res)
        res.extend(url_res)
        return res

    # 处理PDF文件
    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        # 获取布局识别器配置
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")
        
        # 如果需要分析超链接且是根调用
        if parser_config.get("analyze_hyperlink", False) and is_root:
            urls = extract_links_from_pdf(binary)

        # 规范化布局识别器配置
        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"

        # 选择合适的解析器
        name = layout_recognizer.strip().lower()
        parser = PARSERS.get(name, by_plaintext)
        callback(0.1, "Start to parse.")  # 报告解析开始

        # 调用解析器处理PDF文件
        sections, tables, pdf_parser = parser(
            filename = filename,
            binary = binary,
            from_page = from_page,
            to_page = to_page,
            lang = lang,
            callback = callback,
            layout_recognizer = layout_recognizer,
            **kwargs
        )

        # 如果没有解析到内容则返回空列表
        if not sections and not tables:
            return []

        # 某些解析器需要特殊处理
        if name in ["tcadp", "docling", "mineru"]:
            parser_config["chunk_token_num"] = 0

        # 处理表格内容
        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")  # 报告解析完成

    # 处理Excel和CSV文件
    elif re.search(r"\.(csv|xlsx?)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")  # 报告解析开始

        # 检查是否使用TCADP解析器
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")
        if layout_recognizer == "TCADP Parser":
            # 配置TCADP解析器参数
            table_result_type = parser_config.get("table_result_type", "1")
            markdown_image_response_type = parser_config.get("markdown_image_response_type", "1")
            tcadp_parser = TCADPParser(
                table_result_type=table_result_type,
                markdown_image_response_type=markdown_image_response_type
            )
            
            # 检查TCADP解析器是否可用
            if not tcadp_parser.check_installation():
                callback(-1, "TCADP parser not available. Please check Tencent Cloud API configuration.")
                return res

            # 根据文件扩展名确定文件类型
            file_type = "XLSX" if re.search(r"\.xlsx?$", filename, re.IGNORECASE) else "CSV"

            # 使用TCADP解析器处理文件
            sections, tables = tcadp_parser.parse_pdf(
                filepath=filename,
                binary=binary,
                callback=callback,
                output_dir=os.environ.get("TCADP_OUTPUT_DIR", ""),
                file_type=file_type
            )
            
            parser_config["chunk_token_num"] = 0
            res = tokenize_table(tables, doc, is_english)
            callback(0.8, "Finish parsing.")  # 报告解析完成
        else:
            # 使用默认的DeepDOC解析器
            excel_parser = ExcelParser()
            if parser_config.get("html4excel"):
                sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
                parser_config["chunk_token_num"] = 0
            else:
                sections = [(_, "") for _ in excel_parser(binary) if _]

    # 处理文本文件（包括各种编程语言源文件）
    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")  # 报告解析开始
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", 128),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")  # 报告解析完成

    # 处理Markdown文件
    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")  # 报告解析开始
        markdown_parser = Markdown(int(parser_config.get("chunk_token_num", 128)))
        sections, tables, section_images = markdown_parser(
            filename,
            binary,
            separate_tables=False,
            delimiter=parser_config.get("delimiter", "\n!?;。；！？"),
            return_section_images=True,  # 返回段落关联的图像
        )

        final_sections = True  # 标记使用最终格式化的段落

        # 尝试初始化视觉模型用于增强图像识别
        vision_model = None



        # 提取超链接
        if parser_config.get("hyperlink_urls", False) and is_root:
            for idx, (section_text, _) in enumerate(sections):
                soup = markdown_parser.md_to_html(section_text)
                hyperlink_urls = markdown_parser.get_hyperlink_urls(soup)
                urls.update(hyperlink_urls)
                
        # 处理表格内容
        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")  # 报告解析完成

    # 处理HTML文件
    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")  # 报告解析开始
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = HtmlParser()(filename, binary, chunk_token_num)
        sections = [(_, "") for _ in sections if _]  # 过滤空内容
        callback(0.8, "Finish parsing.")  # 报告解析完成

    # 处理JSON文件
    elif re.search(r"\.(json|jsonl|ldjson)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")  # 报告解析开始
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]  # 过滤空内容
        callback(0.8, "Finish parsing.")  # 报告解析完成

    # 处理DOC文件（旧版Word文档）
    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")  # 报告解析开始

        # 尝试导入tika解析器
        try:
            from tika import parser as tika_parser
        except Exception as e:
            callback(0.8, f"tika not available: {e}. Unsupported .doc parsing.")
            logging.warning(f"tika not available: {e}. Unsupported .doc parsing for {filename}.")
            return []

        # 使用tika解析DOC文件
        binary = BytesIO(binary)
        doc_parsed = tika_parser.from_buffer(binary)
        if doc_parsed.get('content', None) is not None:
            sections = doc_parsed['content'].split('\n')
            sections = [(_, "") for _ in sections if _]  # 过滤空行
            callback(0.8, "Finish parsing.")  # 报告解析完成
        else:
            callback(0.8, f"tika.parser got empty content from {filename}.")
            logging.warning(f"tika.parser got empty content from {filename}.")
            return []
    else:
        # 不支持的文件类型
        raise NotImplementedError(
            f"file type not supported yet(pdf, xlsx, doc, docx, txt supported) file: {filename}")

    # 记录合并开始时间
    st = timer()
    
    # 处理最终格式化的段落（主要用于Markdown）
    if final_sections:
        merged_chunks = []  # 存储合并后的文本块
        merged_images = []  # 存储合并后的图像
        chunk_limit = max(0, int(parser_config.get("chunk_token_num", 128)))
        overlapped_percent = int(parser_config.get("overlapped_percent", 0))
        overlapped_percent = max(0, min(overlapped_percent, 90))  # 限制重叠百分比范围

        current_text = ""  # 当前合并块的文本
        current_tokens = 0  # 当前合并块的token数量
        current_image = None  # 当前合并块的图像

        # 逐段合并
        for idx, sec in enumerate(sections):
            text = sec[0] if isinstance(sec, tuple) else sec
            sec_tokens = num_tokens_from_string(text)  # 计算段落token数
            sec_image = section_images[idx] if section_images and idx < len(section_images) else None

            # 如果添加当前段落会超过token限制，则保存当前块并重置
            if current_text and current_tokens + sec_tokens > chunk_limit:
                merged_chunks.append(current_text)
                merged_images.append(current_image)
                overlap_part = ""
                # 处理重叠部分
                if overlapped_percent > 0:
                    overlap_len = int(len(current_text) * overlapped_percent / 100)
                    if overlap_len > 0:
                        overlap_part = current_text[-overlap_len:]
                current_text = overlap_part
                current_tokens = num_tokens_from_string(current_text)
                current_image = current_image if overlap_part else None

            # 添加当前段落文本
            if current_text:
                current_text += "\n" + text
            else:
                current_text = text
            current_tokens += sec_tokens

            # 合并图像
            if sec_image:
                current_image = concat_img(current_image, sec_image) if current_image else sec_image

        # 保存最后一个合并块
        if current_text:
            merged_chunks.append(current_text)
            merged_images.append(current_image)

        chunks = merged_chunks
        has_images = merged_images and any(img is not None for img in merged_images)
        
        # 如果只需返回段落级别分块
        if kwargs.get("section_only", False):
            chunks.extend(embed_res)
            return chunks
        
        # 根据是否有图像选择不同的分词处理方法
        if has_images:
            res.extend(tokenize_chunks_with_images(chunks, doc, is_english, merged_images))
        else:
            res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))
    else:
        # 处理非最终格式化的段落
        # 清理空图像引用
        if section_images:
            if all(image is None for image in section_images):
                section_images = None

        # 带图像的合并
        if section_images:
            chunks, images = naive_merge_with_images(sections, section_images,
                                            int(parser_config.get("chunk_token_num", 128)), 
                                            parser_config.get("delimiter", "\n!?。；！？"))
            
            # 如果只需返回段落级别分块
            if kwargs.get("section_only", False):
                chunks.extend(embed_res)
                return chunks

            res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images))
        else:
            # 不带图像的合并
            chunks = naive_merge(
                sections, int(parser_config.get("chunk_token_num", 128)), 
                parser_config.get("delimiter", "\n!?。；！？"))
            
            # 如果只需返回段落级别分块
            if kwargs.get("section_only", False):
                chunks.extend(embed_res)
                return chunks

            res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))

    # 处理超链接内容
    if urls and parser_config.get("analyze_hyperlink", False) and is_root:
        for index, url in enumerate(urls):
            html_bytes, metadata = extract_html(url)
            if not html_bytes:
                continue
            try:
                sub_url_res = chunk(url, html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
            except Exception as e:
                logging.info(f"Failed to chunk url in registered file type {url}: {e}")
                sub_url_res = chunk(f"{index}.html", html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
            url_res.extend(sub_url_res)

    # 记录合并耗时
    #logging.info("naive_merge({}): {}".format(filename, timer() - st))

    # 添加嵌入文件和超链接处理结果
    if embed_res:
        res.extend(embed_res)
    if url_res:
        res.extend(url_res)
    
    return res

if __name__ == "__main__":
    import sys
    from common.file_utils import get_project_base_directory
    def dummy(prog=None, msg=""):
        pass
    path = os.path.join(get_project_base_directory(), "static/测试文件/test.html")
    result = chunk(path, from_page=0, to_page=100, callback=dummy)
    pass