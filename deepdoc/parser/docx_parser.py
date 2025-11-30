#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import sys, os; __name__ == "__main__" and sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from common.file_utils import get_project_base_directory
from docx import Document
import re
import pandas as pd
from collections import Counter
from rag.nlp import rag_tokenizer
from io import BytesIO


class RAGFlowDocxParser:

    def __extract_table_content(self, tb):
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):

        def blockType(b):
            pattern = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in pattern:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not necessarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        """
        解析DOCX文档并提取指定页码范围内的文本内容和表格数据
        
        该方法是类的可调用接口，支持从文件路径或字节流中加载DOCX文档，
        并根据指定的页码范围提取文档中的段落内容和表格数据。
        段落内容会与对应的样式名称一起保存，表格数据会经过特殊处理。
        
        参数:
            fnm: str或bytes - DOCX文档的文件路径或字节流数据
            from_page: int - 起始页码（包含），默认为0
            to_page: int - 结束页码（不包含），默认为一个大数，以确保解析所有页面
            
        返回:
            tuple - 包含两个元素的元组：
                1. secs: list - 解析出的段落内容列表，每个元素为(文本内容, 样式名称)的元组
                2. tbls: list - 解析出的表格内容列表，每个表格经过__extract_table_content处理
        """
        # 初始化文档对象，根据输入类型选择不同的加载方式
        # 如果是字符串，认为是文件路径；否则作为字节流处理
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0  # 已解析的当前页码
        secs = []  # 存储解析出的段落内容
        
        # 遍历文档中的所有段落
        for p in self.doc.paragraphs:
            if pn > to_page:  # 如果已超过指定的结束页码，提前退出
                break

            # 保存单个段落中在指定页码范围内的所有run文本
            runs_within_single_paragraph = []
            for run in p.runs:
                if pn > to_page:  # 如果已超过指定的结束页码，提前退出run循环
                    break  
                # 如果当前页码在指定范围内，且段落文本不为空，则保存run的文本
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)  # 先收集run.text

                # 检查是否存在分页符标记，如有则页码加1
                # 注：可以考虑将此逻辑封装为静态方法以提高代码可读性
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1

            # 将收集到的run文本合并，并与段落样式一起添加到结果列表
            # 安全地获取样式名称，避免属性不存在的错误
            secs.append("""join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else ''""") 

        # 提取文档中所有表格的内容
        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls

if __name__ == "__main__":

    parser = RAGFlowDocxParser()
    path = os.path.join(get_project_base_directory(), "static/测试文件/test.docx")
    secs, tbls = parser(path)
    print(secs)
    print(tbls)